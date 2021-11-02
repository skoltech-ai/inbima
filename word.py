import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Mm, Cm, Pt


class Word():
    def __init__(self, years, get_papers):
        self.years = years
        self.get_papers = get_papers

        self.document = docx.Document()

        self.style_normal = self.document.styles['Normal']
        self.style_normal.font.name = 'Cambria'
        self.style_normal.font.size = Pt(12)

        self.style_title_name = self.document.styles.add_style('TitleName',
            WD_STYLE_TYPE.PARAGRAPH)
        self.style_title_name.base_style = self.style_normal
        self.style_title_name.font.name = 'Bookman Old Style'
        self.style_title_name.font.size = Pt(20)

        self.table_stat_title = self.document.styles.add_style('TableStatTitle',
            WD_STYLE_TYPE.PARAGRAPH)
        self.table_stat_title.base_style = self.style_normal
        self.table_stat_title.font.name = 'Bookman Old Style'
        self.table_stat_title.font.size = Pt(12)

    def add_break(self):
        self.document.add_page_break()

    def add_grant_info(self, grant, head, photo_logo=None):
        table = self.document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].width = Cm(6.5)
        table.rows[0].cells[0].height = Cm(20)
        table.rows[0].cells[1].width = Cm(12)
        table.rows[0].cells[1].height = Cm(20)

        table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = table.rows[0].cells[0].paragraphs[0]
        if photo_logo:
            p.add_run().add_picture(photo_logo, width=Cm(6))

        table = table.rows[0].cells[1].add_table(rows=5, cols=2)
        table.rows[0].cells[0].width = Cm(3)
        table.rows[0].cells[1].width = Cm(9)

        table.rows[0].cells[0].text = 'Grant id'
        table.rows[0].cells[1].text = str(grant.get('id', ''))

        table.rows[1].cells[0].text = 'Head of the project'
        head = head['name'] + ' ' + head['surname']
        table.rows[1].cells[1].text = head

        table.rows[2].cells[0].text = 'Grant Number'
        table.rows[2].cells[1].text = grant.get('number', '')

        table.rows[3].cells[0].text = 'Source'
        table.rows[3].cells[1].text = grant.get('source', '')

        table.rows[4].cells[0].text = 'Acknowledgement'
        table.rows[4].cells[1].text = grant.get('acknowledgement', '')

    def add_note(self, is_grant=False):
        self.document.add_paragraph('\n\n')
        text = 'This document is generated automatically. '
        text += 'Journal ratings were determined based on the '
        text += 'Scimago Journal & Country Rank. All publications '
        text += 'in the collected database will be manually checked later. '
        if is_grant:
            text += '\n\nNOTE! We use bold type for all authors associated with our center. This will be clarified later (only authors participating in the corresponding grant will be displayed in bold).'
        p = self.document.add_paragraph(text, style='Intense Quote')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_paper_list(self, stat, author=None, grant=None, with_links=False):
        self.document.add_heading('Publications', level=1)

        for q in [1, 2, 0]:
            if stat['total'][f'q{q}'] == 0:
                continue
            if q == 1:
                text = 'Publications in Q1-rated journals'
            elif q == 2:
                text = 'Publications in Q2-rated journals'
            else:
                text = 'Other publications'
            self.document.add_heading(text, level=3)

            for year in self.years[::-1]:
                if stat[year][f'q{q}'] == 0:
                    continue

                papers = self.get_papers(author, year, q, grant)

                self.document.add_heading('Year ' + str(year), level=5)

                for paper in papers.values():
                    journal = paper['journal_object']

                    p = self.document.add_paragraph(style='List Bullet')
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    authors = paper['authors'].split(', ')
                    authors_parsed = paper['authors_parsed'].split(', ')

                    for i in range(len(authors)):
                        text = p.add_run(authors[i])
                        if authors_parsed[i] in (author or ''):
                            text.bold = True
                        if not author and authors_parsed[i][0] == '#':
                            text.bold = True
                        if i < len(authors)-1:
                            p.add_run(', ')
                        else:
                            p.add_run('. ')

                    p.add_run(paper['year'] + '. ')
                    p.add_run(paper['title'] + '. ')
                    p.add_run(paper['journal'] + '. ').italic = True
                    p.add_run(compose_paper_nums(paper))
                    if with_links:
                        v = paper.get('screen')
                        if v and len(v) > 2:
                            p.add_run(' // ')
                            add_hyperlink(p, 'Screenshot Publisher', v)
                        v = journal.get('screen_wos')
                        if v and len(v) > 2:
                            p.add_run(' // ')
                            add_hyperlink(p, 'Screenshot WoS', v)

            p.add_run('\n\n')

    def add_person_info(self, person, photo_person=None, photo_logo=None):
        table = self.document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].width = Cm(7.5)
        table.rows[0].cells[0].height = Cm(20)
        table.rows[0].cells[1].width = Cm(11)
        table.rows[0].cells[1].height = Cm(20)

        p = table.rows[0].cells[0].paragraphs[0]

        if photo_person:
            p.add_run().add_picture(photo_person, width=Cm(7))

        p = table.rows[0].cells[1].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if photo_logo:
            p.add_run().add_picture(photo_logo, width=Cm(5))

        p = table.rows[0].cells[1].add_paragraph()
        p.style = self.style_title_name
        name_full = person.get('name', '') + ' ' + person.get('surname', '')
        p.add_run('\n' + name_full).bold = True

        table = table.rows[0].cells[1].add_table(rows=7, cols=2)
        table.rows[0].cells[0].width = Cm(3)
        table.rows[0].cells[1].width = Cm(8)

        # table.rows[0].cells[0].text = 'Year of birth'
        # table.rows[0].cells[1].text = str(person.get('birth', ''))

        table.rows[1].cells[0].text = 'Degree'
        table.rows[1].cells[1].text = person.get('degree', '')

        table.rows[2].cells[0].text = 'Position'
        table.rows[2].cells[1].text = person.get('position', '')

        v = person.get('page')
        if v:
            table.rows[3].cells[0].text = 'Personal page'
            p = table.rows[3].cells[1].paragraphs[0]
            add_hyperlink(p, 'Skoltech profile', v)

        v = person.get('h_sch')
        if v:
            table.rows[4].cells[0].text = 'H-index Scholar'
            p = table.rows[4].cells[1].paragraphs[0]
            p.add_run(str(v)).bold = True
            p.add_run('\t[')
            add_hyperlink(p, 'link', person.get('link_sch'))
            p.add_run(']')

        v = person.get('h_sco')
        if v:
            table.rows[5].cells[0].text = 'H-index Scopus'
            p = table.rows[5].cells[1].paragraphs[0]
            p.add_run(str(v)).bold = True
            p.add_run('\t[')
            add_hyperlink(p, 'link', person.get('link_sco'))
            p.add_run(']')

        v = person.get('h_wos')
        if v:
            table.rows[6].cells[0].text = 'H-index WoS'
            p = table.rows[6].cells[1].paragraphs[0]
            p.add_run(str(v)).bold = True
            p.add_run('\t[')
            add_hyperlink(p, 'link', person.get('link_wos'))
            p.add_run(']')

    def add_person_stat(self, stat):
        self.document.add_paragraph('\n\n')
        table = self.document.add_table(rows=5, cols=7)
        table.rows[0].cells[0].width = Cm(5)
        table.style = 'Table Grid'

        for i, year in enumerate(self.years, 1):
            p = table.rows[0].cells[i].paragraphs[0]
            p.style = self.table_stat_title
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(str(year))
        p = table.rows[0].cells[6].paragraphs[0]
        p.style = self.table_stat_title
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('Total').bold = True

        p = table.rows[1].cells[0].paragraphs[0]
        p.style = self.table_stat_title
        p.add_run('Q1 journals')
        p = table.rows[2].cells[0].paragraphs[0]
        p.style = self.table_stat_title
        p.add_run('Q2 journals')
        p = table.rows[3].cells[0].paragraphs[0]
        p.style = self.table_stat_title
        p.add_run('Other journals')
        p = table.rows[4].cells[0].paragraphs[0]
        p.style = self.table_stat_title
        p.add_run('Total').bold = True

        for i, year in enumerate(self.years, 1):
            p = table.rows[1].cells[i].paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(str(stat[year]['q1']))
            p = table.rows[2].cells[i].paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(str(stat[year]['q2']))
            p = table.rows[3].cells[i].paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(str(stat[year]['q0']))
            p = table.rows[4].cells[i].paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(str(stat[year]['total']))

        p = table.rows[1].cells[6].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(str(stat['total']['q1']))
        p = table.rows[2].cells[6].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(str(stat['total']['q2']))
        p = table.rows[3].cells[6].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(str(stat['total']['q0']))
        p = table.rows[4].cells[6].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(str(stat['total']['total']))

    def save(self, file_path):
        for section in self.document.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(1)
        self.document.save(file_path)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def compose_paper_nums(paper, end='.'):
    volume = paper.get('volume')
    number = paper.get('number')
    pages = paper.get('pages')
    text = ''
    if volume:
        text += f'{volume}({number})' if number else f'{volume}'
        text += f': {pages}' if pages else ''
        text += end
    elif pages:
        text += f'{pages}'
        text += end
    return text
